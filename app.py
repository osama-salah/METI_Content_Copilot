import streamlit as st
import google.generativeai as genai
from io import BytesIO
import re

# --- Optional dependencies for download functionality ---
try:
    from fpdf import FPDF
    import docx
    from docx.shared import Pt
    DOWNLOAD_ENABLED = True
except ImportError:
    DOWNLOAD_ENABLED = False

# Ensure st is available globally before importing helpers
try:
    from helpers.text_utils import extract_text_from_files
    from helpers import prompt_utils
except ImportError:
    # Creating dummy functions if helpers are not available
    # This allows the app to run without the helper files for styling purposes.
    st.warning("Helper modules (helpers.text_utils, helpers.prompt_utils) not found. Using dummy functions. App functionality will be limited.")
    def extract_text_from_files(files):
        return " ".join([file.name for file in files])
    class PromptUtils:
        def create_generation_prompt(self, *args): return "Dummy generation prompt"
        def create_validation_prompt(self, *args): return "Dummy validation prompt"
        def create_updater_prompt(self, *args): return "Dummy updater prompt"
        def create_quiz_creator_prompt(self, *args): return "Dummy quiz prompt"
    prompt_utils = PromptUtils()

except NameError as e:
    st.error(f"Import error in helper modules: {e}")
    st.error("Make sure all helper modules that use 'st' have 'import streamlit as st' at the top")
    st.stop()
except Exception as e:
    st.error(f"Unexpected import error: {e}")
    st.stop()

# --- Helper functions for file creation with Markdown parsing ---
def create_descriptive_filename(content_type, user_params=None, content_text="", file_extension="pdf"):
    """Create descriptive filenames based on content and user selections"""
    from datetime import datetime
    
    # Extract course topic/title if possible
    topic = "Course"
    if content_text:
        lines = content_text.split('\n')
        for line in lines:
            if line.strip().startswith('# '):
                topic = line.strip()[2:].strip()
                # Clean for filename
                topic = re.sub(r'[^\w\s-]', '', topic)
                topic = re.sub(r'[-\s]+', '_', topic)[:30]
                break
    
    # Build filename parts
    parts = [topic]
    
    if user_params:
        if 'audience' in user_params:
            audience_short = user_params['audience'].split()[0]  # "High", "Undergraduate", etc.
            parts.append(audience_short)
        
        if 'length' in user_params:
            length_short = user_params['length'].split()[0]  # "Quick", "Moderate", "Detailed"
            parts.append(length_short)
            
        if 'quiz_type' in user_params:
            quiz_type_short = user_params['quiz_type'].replace(' ', '')
            parts.append(quiz_type_short)
    
    parts.append(content_type)
    
    # Add timestamp for uniqueness
    timestamp = datetime.now().strftime("%m%d_%H%M")
    parts.append(timestamp)
    
    filename = "_".join(parts) + f".{file_extension}"
    return filename
    
def create_styled_docx(text):
    """Generates a DOCX file from a Markdown string with styling for headers, lists, and code."""
    doc = docx.Document()
    in_code_block = False
    code_block_started = False  # Track if we just started a code block
    
    for line in text.split('\n'):
        # Code block handling - FIX: Check for triple backticks
        if line.strip().startswith('```'):
            in_code_block = not in_code_block
            if in_code_block:
                # Only add spacing before code block if the previous element wasn't empty
                code_block_started = True
            else:
                # Add some spacing after code block
                doc.add_paragraph()
                code_block_started = False
            continue

        if in_code_block:
            # For code blocks, add paragraph with shaded background
            p = doc.add_paragraph(line)
            p.style = 'No Spacing'
            
            # Set font for code
            if p.runs:
                font = p.runs[0].font
            else:
                run = p.add_run('')
                font = run.font
            font.name = 'Courier New'
            font.size = Pt(10)
            
            # Add shaded background to the paragraph
            from docx.oxml.shared import qn
            from docx.oxml import parse_xml
            
            # Create shading element with light gray background
            shading_elm = parse_xml(r'<w:shd {} w:fill="F0F0F0"/>'.format(
                'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            ))
            p._element.get_or_add_pPr().append(shading_elm)
            
            # Add border around code block paragraphs
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
            # Set paragraph spacing for code blocks
            paragraph_format = p.paragraph_format
            paragraph_format.left_indent = Pt(12)  # Slight indentation
            paragraph_format.right_indent = Pt(12)
            
            # Only add space before the first line of the code block
            if code_block_started:
                paragraph_format.space_before = Pt(6)
                code_block_started = False
            else:
                paragraph_format.space_before = Pt(2)
            
            paragraph_format.space_after = Pt(2)
            
            continue

        # Other markdown handling (rest of the function remains the same)
        line = line.strip()
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith(('* ', '- ')):
            p = doc.add_paragraph(style='List Bullet')
            # Handle inline styles (bold, italic, code)
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|\`.*?\`)', line[2:])
            for part in parts:
                if not part: continue
                if part.startswith('**') and part.endswith('**'):
                    p.add_run(part[2:-2]).bold = True
                elif part.startswith('*') and part.endswith('*'):
                    p.add_run(part[1:-1]).italic = True
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Courier New'
                    # Add light gray background for inline code
                    from docx.oxml.shared import qn
                    from docx.oxml import parse_xml
                    shading_elm = parse_xml(r'<w:shd {} w:fill="E8E8E8"/>'.format(
                        'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
                    ))
                    run._element.get_or_add_rPr().append(shading_elm)
                else:
                    p.add_run(part)
        elif line:
            p = doc.add_paragraph()
            # Handle inline styles (bold, italic, code)
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|\`.*?\`)', line)
            for part in parts:
                if not part: continue
                if part.startswith('**') and part.endswith('**'):
                    p.add_run(part[2:-2]).bold = True
                elif part.startswith('*') and part.endswith('*'):
                    p.add_run(part[1:-1]).italic = True
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Courier New'
                    # Add light gray background for inline code
                    from docx.oxml.shared import qn
                    from docx.oxml import parse_xml
                    shading_elm = parse_xml(r'<w:shd {} w:fill="E8E8E8"/>'.format(
                        'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
                    ))
                    run._element.get_or_add_rPr().append(shading_elm)
                else:
                    p.add_run(part)

    doc_fp = BytesIO()
    doc.save(doc_fp)
    doc_fp.seek(0)
    return doc_fp.getvalue()

def create_html_docx(html_text):
    """
    Generates a DOCX file from an HTML string with styling 
    for headers, lists, bold, italic, and other HTML elements.
    """
    from html.parser import HTMLParser
    import html
    
    class HTMLToDocxParser(HTMLParser):
        def __init__(self, doc):
            super().__init__()
            self.doc = doc
            self.current_paragraph = None
            self.current_run = None
            self.style_stack = []  # Stack to track nested styles
            self.in_list = False
            self.list_stack = []  # Stack to track nested lists
            self.list_counters = []  # Stack to track list item counters
            self.current_text = ""
            self.pending_paragraph = False
            
        def handle_starttag(self, tag, attrs):
            self.flush_text()
            
            if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(tag[1])
                # Add spacing before headers (except if it's the first element)
                if len(self.doc.paragraphs) > 0:
                    self.doc.add_paragraph()
                self.current_paragraph = self.doc.add_heading('', level=level)
                self.current_run = None
            elif tag == 'p':
                # Only add paragraph if we don't already have one or if we have content
                if self.current_paragraph is None or len(self.doc.paragraphs) > 0:
                    self.current_paragraph = self.doc.add_paragraph()
                self.current_run = None
            elif tag in ['b', 'strong']:
                self.style_stack.append('bold')
                self._ensure_run()
            elif tag in ['i', 'em']:
                self.style_stack.append('italic')
                self._ensure_run()
            elif tag == 'code':
                self.style_stack.append('code')
                self._ensure_run()
            elif tag == 'pre':
                # Add spacing before pre block
                if len(self.doc.paragraphs) > 0:
                    self.doc.add_paragraph()
                self.current_paragraph = self.doc.add_paragraph()
                self.current_paragraph.style = 'No Spacing'
                self.style_stack.append('pre')
                self._ensure_run()
            elif tag == 'ul':
                self.in_list = True
                self.list_stack.append('bullet')
                self.list_counters.append(0)  # Not used for bullets but keeps stack aligned
                # Add spacing before list
                if len(self.doc.paragraphs) > 0 and not self.current_paragraph:
                    self.doc.add_paragraph()
            elif tag == 'ol':
                self.in_list = True
                self.list_stack.append('number')
                self.list_counters.append(0)  # Counter for numbered items
                # Add spacing before list
                if len(self.doc.paragraphs) > 0 and not self.current_paragraph:
                    self.doc.add_paragraph()
            elif tag == 'li':
                if self.list_stack:
                    list_type = self.list_stack[-1]
                    if list_type == 'number':
                        # Increment counter for this list level
                        self.list_counters[-1] += 1
                        self.current_paragraph = self.doc.add_paragraph(style='List Number')
                    else:
                        self.current_paragraph = self.doc.add_paragraph(style='List Bullet')
                else:
                    # Fallback to bullet if no list context
                    self.current_paragraph = self.doc.add_paragraph(style='List Bullet')
                self.current_run = None
            elif tag == 'br':
                if self.current_paragraph is None:
                    self.current_paragraph = self.doc.add_paragraph()
                self.current_paragraph.add_run().add_break()
                
        def handle_endtag(self, tag):
            self.flush_text()
            
            if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                self.current_paragraph = None
                self.current_run = None
            elif tag == 'p':
                self.current_paragraph = None
                self.current_run = None
            elif tag in ['b', 'strong', 'i', 'em', 'code']:
                if self.style_stack and self.style_stack[-1] in ['bold', 'italic', 'code']:
                    self.style_stack.pop()
                self.current_run = None  # Force new run creation
            elif tag == 'pre':
                if 'pre' in self.style_stack:
                    self.style_stack.remove('pre')
                self.current_paragraph = None
                self.current_run = None
                # Add spacing after pre block
                self.doc.add_paragraph()
            elif tag in ['ul', 'ol']:
                if self.list_stack:
                    self.list_stack.pop()
                    self.list_counters.pop()
                self.in_list = len(self.list_stack) > 0
                if not self.in_list:
                    # Add spacing after outermost list
                    self.doc.add_paragraph()
            elif tag == 'li':
                self.current_paragraph = None
                self.current_run = None
                
        def handle_data(self, data):
            # Decode HTML entities and clean up excessive whitespace
            data = html.unescape(data)
            # Only add non-empty text
            if data.strip():
                self.current_text += data.strip() + " "
            
        def flush_text(self):
            if self.current_text.strip():
                self._ensure_paragraph()
                self._ensure_run()
                
                # Apply styles from stack
                text_to_add = self.current_text.strip()
                run = self.current_run.add_run(text_to_add) if hasattr(self.current_run, 'add_run') else self.current_paragraph.add_run(text_to_add)
                
                # Apply formatting based on style stack
                if 'bold' in self.style_stack:
                    run.bold = True
                if 'italic' in self.style_stack:
                    run.italic = True
                if 'code' in self.style_stack or 'pre' in self.style_stack:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                    # Add light gray background for code
                    from docx.oxml.shared import qn
                    from docx.oxml import parse_xml
                    try:
                        shading_elm = parse_xml(r'<w:shd {} w:fill="E8E8E8"/>'.format(
                            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
                        ))
                        run._element.get_or_add_rPr().append(shading_elm)
                    except:
                        pass  # Skip styling if it fails
                
                self.current_text = ""
        
        def _ensure_paragraph(self):
            if self.current_paragraph is None:
                self.current_paragraph = self.doc.add_paragraph()
        
        def _ensure_run(self):
            if self.current_run is None:
                self._ensure_paragraph()
                self.current_run = self.current_paragraph
    
    # Create the document and parse the HTML
    doc = docx.Document()
    parser = HTMLToDocxParser(doc)
    parser.feed(html_text)
    parser.flush_text()  # Flush any remaining text
    
    # Save to BytesIO and return
    doc_fp = BytesIO()
    doc.save(doc_fp)
    doc_fp.seek(0)
    return doc_fp.getvalue()

def process_pdf_inline_styles(pdf, line, indent):
    """Helper to process a line with inline markdown for PDF `write` method, handling wrapping."""
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|\`.*?\`)', line)
    for part in parts:
        if not part: continue
        
        style = ''
        font = 'Arial'
        size = 12
        if part.startswith('**') and part.endswith('**'):
            content = part[2:-2]
            style = 'B'
        elif part.startswith('*') and part.endswith('*'):
            content = part[1:-1]
            style = 'I'
        elif part.startswith('`') and part.endswith('`'):
            content = part[1:-1]
            font = 'Courier'
            style = ''
            size = 10
        else:
            content = part
        
        pdf.set_font(font, style, size)
        
        # Manual word wrapping for the `write` method
        words = content.split(' ')
        for word in words:
            word_to_write = word + ' '
            word_width = pdf.get_string_width(word_to_write)
            if pdf.get_x() + word_width > pdf.w - pdf.r_margin:
                pdf.ln()
                pdf.set_x(indent) # Re-apply indent for wrapped line
            pdf.write(5, word_to_write.encode('latin-1', 'replace').decode('latin-1'))
    
    pdf.set_font('Arial', '', 12) # Reset font at the end of the line

def create_styled_pdf(text):
    """
    Generates a PDF file from a Markdown string with native styling
    for headers, lists, bold, italic, and code.
    """
    try:
        pdf = FPDF()
        pdf.add_page()
        
        # Simplified font handling - stick to built-in fonts
        pdf.set_font('Arial', size=12)
        unicode_support = False  # Disable Unicode to avoid font issues
        default_font = 'Arial'
        
        effective_page_width = pdf.w - 2 * pdf.l_margin

        in_code_block = False
        for line in text.split('\n'):
            # Code block handling
            if line.strip().startswith('```'):
                in_code_block = not in_code_block
                if in_code_block:
                    pdf.set_font('Courier', size=10)
                    pdf.set_fill_color(240, 240, 240)
                    pdf.ln(2)
                else:
                    pdf.set_font('Arial', size=12)
                    pdf.set_fill_color(255, 255, 255)
                    pdf.ln(5)
                continue

            if in_code_block:
                pdf.set_x(pdf.l_margin)
                # Ensure safe encoding for code blocks
                safe_line = line.encode('latin-1', 'replace').decode('latin-1')
                pdf.multi_cell(effective_page_width, 5, safe_line, border=0, fill=True)
                continue

            # Other markdown handling
            line = line.strip()
            if not line:
                pdf.ln(5)
                continue

            # Safe text encoding for all content
            if line.startswith('# '):
                pdf.set_font('Arial', 'B', 16)
                safe_text = line[2:].strip().encode('latin-1', 'replace').decode('latin-1')
                pdf.multi_cell(effective_page_width, 8, safe_text)
                pdf.ln(4)
            elif line.startswith('## '):
                pdf.set_font('Arial', 'B', 14)
                safe_text = line[3:].strip().encode('latin-1', 'replace').decode('latin-1')
                pdf.multi_cell(effective_page_width, 7, safe_text)
                pdf.ln(3)
            elif line.startswith('### '):
                pdf.set_font('Arial', 'B', 12)
                safe_text = line[4:].strip().encode('latin-1', 'replace').decode('latin-1')
                pdf.multi_cell(effective_page_width, 6, safe_text)
                pdf.ln(2)
            elif line.startswith(('* ', '- ')):
                pdf.set_x(pdf.l_margin)
                initial_x = pdf.get_x()
                pdf.write(5, '- ')
                process_pdf_inline_styles_safe(pdf, line[2:].strip(), initial_x + 5)
                pdf.ln()
            else:
                pdf.set_x(pdf.l_margin)
                process_pdf_inline_styles_safe(pdf, line, pdf.l_margin)
                pdf.ln()

        pdf_fp = BytesIO()
        pdf.output(pdf_fp)
        pdf_fp.seek(0)
        return pdf_fp.getvalue()
        
    except Exception as e:
        # If PDF generation fails, create a simple fallback PDF
        st.error(f"PDF generation failed: {e}")
        return create_fallback_pdf(text)

def process_pdf_inline_styles_safe(pdf, line, indent):
    """Simplified version of inline style processing with better error handling"""
    try:
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|\`.*?\`)', line)
        for part in parts:
            if not part: 
                continue
            
            style = ''
            font = 'Arial'
            size = 12
            
            if part.startswith('**') and part.endswith('**'):
                content = part[2:-2]
                style = 'B'
            elif part.startswith('*') and part.endswith('*'):
                content = part[1:-1]
                style = 'I'
            elif part.startswith('`') and part.endswith('`'):
                content = part[1:-1]
                font = 'Courier'
                size = 10
            else:
                content = part
            
            # Ensure safe encoding
            safe_content = content.encode('latin-1', 'replace').decode('latin-1')
            
            pdf.set_font(font, style, size)
            
            # Simple word wrapping
            words = safe_content.split(' ')
            for word in words:
                word_to_write = word + ' '
                word_width = pdf.get_string_width(word_to_write)
                if pdf.get_x() + word_width > pdf.w - pdf.r_margin:
                    pdf.ln()
                    pdf.set_x(indent)
                pdf.write(5, word_to_write)
        
        pdf.set_font('Arial', '', 12)
        
    except Exception as e:
        # Fallback to simple text if styling fails
        safe_text = line.encode('latin-1', 'replace').decode('latin-1')
        pdf.set_font('Arial', '', 12)
        pdf.multi_cell(pdf.w - 2 * pdf.l_margin, 5, safe_text)

def create_fallback_pdf(text):
    """Create a simple PDF if the styled version fails"""
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font('Arial', size=12)
        
        # Simple text conversion - remove markdown
        clean_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)  # Remove bold
        clean_text = re.sub(r'\*([^*]+)\*', r'\1', clean_text)  # Remove italic
        clean_text = re.sub(r'`([^`]+)`', r'\1', clean_text)  # Remove code
        clean_text = re.sub(r'^#+\s*', '', clean_text, flags=re.MULTILINE)  # Remove headers
        
        # Encode safely
        safe_text = clean_text.encode('latin-1', 'replace').decode('latin-1')
        
        pdf.multi_cell(pdf.w - 2 * pdf.l_margin, 5, safe_text)
        
        pdf_fp = BytesIO()
        pdf.output(pdf_fp)
        pdf_fp.seek(0)
        return pdf_fp.getvalue()
        
    except Exception as e:
        st.error(f"Even fallback PDF generation failed: {e}")
        return b"PDF generation failed"


def create_html_pdf(html_text):
    """
    Generates a PDF file from an HTML string with styling
    for headers, lists, bold, italic, and other HTML elements.
    Properly handles ordered lists with numbers and nested lists with letters.
    """
    from html.parser import HTMLParser
    import html
    
    class HTMLToPDFParser(HTMLParser):
        def __init__(self, pdf):
            super().__init__()
            self.pdf = pdf
            self.font_stack = [('Arial', '', 12)]  # Stack to track font changes
            self.in_code = False
            self.in_pre = False
            self.list_stack = []  # Stack to track list types and levels
            self.list_counters = []  # Stack to track counters for each list level
            self.list_indent_level = 0
            self.base_indent = 10  # Base indentation for lists
            self.current_text = ""
            self.unicode_support = False
            self.last_was_block_element = False
            self.in_list_item = False
            self.list_item_first_line = True
            
        def get_current_indent(self):
            """Calculate current indentation based on list nesting level"""
            return self.pdf.l_margin + (self.list_indent_level * self.base_indent)
            
        def handle_starttag(self, tag, attrs):
            self.flush_text()
            
            if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                # Add spacing before headers
                if not self.last_was_block_element and self.pdf.get_y() > 20:
                    self.pdf.ln(8)
                
                level = int(tag[1])
                size = max(16 - (level - 1) * 2, 10)
                self.pdf.set_font('Arial', 'B', size)
                self.last_was_block_element = True
                
            elif tag == 'p':
                # Add spacing before paragraphs (but not if we're in a list item)
                if not self.last_was_block_element and self.pdf.get_y() > 20 and not self.in_list_item:
                    self.pdf.ln(4)
                self.last_was_block_element = True
                
            elif tag == 'b' or tag == 'strong':
                current_font = self.font_stack[-1]
                new_style = current_font[1] + 'B' if 'B' not in current_font[1] else current_font[1]
                self.font_stack.append((current_font[0], new_style, current_font[2]))
                self.pdf.set_font(current_font[0], new_style, current_font[2])
                
            elif tag == 'i' or tag == 'em':
                current_font = self.font_stack[-1]
                new_style = current_font[1] + 'I' if 'I' not in current_font[1] else current_font[1]
                self.font_stack.append((current_font[0], new_style, current_font[2]))
                self.pdf.set_font(current_font[0], new_style, current_font[2])
                
            elif tag == 'code':
                self.in_code = True
                self.font_stack.append(('Courier', '', 10))
                self.pdf.set_font('Courier', '', 10)
                
            elif tag == 'pre':
                # Add spacing before pre block
                if not self.last_was_block_element:
                    self.pdf.ln(4)
                self.in_pre = True
                self.pdf.set_font('Courier', '', 10)
                self.pdf.set_fill_color(240, 240, 240)
                self.last_was_block_element = True
                
            elif tag == 'ul':
                # Add spacing before lists (only for outermost lists)
                if self.list_indent_level == 0 and not self.last_was_block_element:
                    self.pdf.ln(4)
                
                self.list_stack.append('ul')
                self.list_counters.append(0)  # Not used for bullets but keeps stack aligned
                self.list_indent_level += 1
                self.last_was_block_element = True
                
            elif tag == 'ol':
                # Add spacing before lists (only for outermost lists)
                if self.list_indent_level == 0 and not self.last_was_block_element:
                    self.pdf.ln(4)
                
                self.list_stack.append('ol')
                self.list_counters.append(0)  # Counter for numbered items
                self.list_indent_level += 1
                self.last_was_block_element = True
                
            elif tag == 'li':
                self.in_list_item = True
                self.list_item_first_line = True
                
                # Move to new line for list item
                if self.pdf.get_x() > self.pdf.l_margin:
                    self.pdf.ln()
                
                # Set proper indentation
                current_indent = self.get_current_indent()
                self.pdf.set_x(current_indent)
                
                if self.list_stack:
                    list_type = self.list_stack[-1]
                    if list_type == 'ol':
                        # Increment counter and use number
                        self.list_counters[-1] += 1
                        counter = self.list_counters[-1]
                        
                        # Format number based on nesting level
                        if self.list_indent_level == 1:
                            # Top level: 1. 2. 3.
                            number_text = f"{counter}. "
                        elif self.list_indent_level == 2:
                            # Second level: a) b) c)
                            letter = chr(ord('a') + counter - 1) if counter <= 26 else f"item{counter}"
                            number_text = f"{letter}) "
                        elif self.list_indent_level == 3:
                            # Third level: i. ii. iii.
                            roman_numerals = ['i', 'ii', 'iii', 'iv', 'v', 'vi', 'vii', 'viii', 'ix', 'x']
                            roman = roman_numerals[counter - 1] if counter <= len(roman_numerals) else f"item{counter}"
                            number_text = f"{roman}. "
                        else:
                            # Deeper levels: fallback to numbers
                            number_text = f"{counter}. "
                            
                        try:
                            self.pdf.write(5, number_text)
                        except:
                            # Fallback for encoding issues
                            safe_text = number_text.encode('ascii', 'replace').decode('ascii')
                            self.pdf.write(5, safe_text)
                    else:
                        # Use different bullet styles based on nesting level
                        if self.list_indent_level == 1:
                            bullet_char = '• ' if self.unicode_support else '* '
                        elif self.list_indent_level == 2:
                            bullet_char = '◦ ' if self.unicode_support else '- '
                        else:
                            bullet_char = '▪ ' if self.unicode_support else '+ '
                            
                        try:
                            self.pdf.write(5, bullet_char)
                        except:
                            self.pdf.write(5, '- ')
                else:
                    # Fallback bullet if no list context
                    try:
                        self.pdf.write(5, '• ')
                    except:
                        self.pdf.write(5, '- ')
                    
            elif tag == 'br':
                self.pdf.ln()
                # Maintain indentation if we're in a list item
                if self.in_list_item:
                    # Calculate the indent for continuation lines (after the bullet/number)
                    current_indent = self.get_current_indent()
                    bullet_width = 15  # Approximate width of bullet/number + space
                    self.pdf.set_x(current_indent + bullet_width)
                    self.list_item_first_line = False
                else:
                    self.pdf.set_x(self.pdf.l_margin)
                
        def handle_endtag(self, tag):
            self.flush_text()
            
            if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                self.pdf.ln(6)  # Space after headers
                self.pdf.set_font('Arial', '', 12)  # Reset font
                self.pdf.set_x(self.pdf.l_margin)  # Reset position
                self.last_was_block_element = True
                
            elif tag == 'p':
                # Only add line break if we're not in a list item
                if not self.in_list_item:
                    self.pdf.ln(4)  # Space after paragraphs
                    self.pdf.set_x(self.pdf.l_margin)  # Reset position
                self.last_was_block_element = True
                
            elif tag in ['b', 'strong', 'i', 'em', 'code']:
                if len(self.font_stack) > 1:
                    self.font_stack.pop()
                current_font = self.font_stack[-1]
                self.pdf.set_font(current_font[0], current_font[1], current_font[2])
                if tag == 'code':
                    self.in_code = False
                    
            elif tag == 'pre':
                self.in_pre = False
                self.pdf.set_fill_color(255, 255, 255)  # Reset background
                self.pdf.set_font('Arial', '', 12)  # Reset font
                self.pdf.ln(6)  # Space after pre block
                self.pdf.set_x(self.pdf.l_margin)  # Reset position
                self.last_was_block_element = True
                
            elif tag in ['ul', 'ol']:
                if self.list_stack:
                    self.list_stack.pop()
                    self.list_counters.pop()
                    self.list_indent_level -= 1
                    
                # Add spacing after outermost list
                if self.list_indent_level == 0:
                    self.pdf.ln(4)
                    self.pdf.set_x(self.pdf.l_margin)
                    self.last_was_block_element = True
                
            elif tag == 'li':
                self.in_list_item = False
                self.list_item_first_line = True
                self.pdf.ln()  # New line after list item
                
        def handle_data(self, data):
            """Handle text data between HTML tags"""
            # Decode HTML entities and clean up excessive whitespace
            data = html.unescape(data)
            # Only add non-empty text
            if data.strip():
                self.current_text += data.strip() + " "
        
        def flush_text(self):
            """Write accumulated text to PDF with proper wrapping and indentation"""
            if self.current_text.strip():
                # Handle Unicode encoding
                safe_text = self.current_text.strip()
                if not self.unicode_support:
                    safe_text = safe_text.encode('latin-1', 'replace').decode('latin-1')
                
                # Calculate available width for text
                if self.in_list_item:
                    current_indent = self.get_current_indent()
                    if self.list_item_first_line:
                        # First line: text starts after bullet/number
                        bullet_width = 15
                        available_width = self.pdf.w - current_indent - bullet_width - self.pdf.r_margin
                        start_x = current_indent + bullet_width
                    else:
                        # Continuation lines: align with first line text
                        bullet_width = 15
                        available_width = self.pdf.w - current_indent - bullet_width - self.pdf.r_margin
                        start_x = current_indent + bullet_width
                        self.pdf.set_x(start_x)
                else:
                    available_width = self.pdf.w - self.pdf.l_margin - self.pdf.r_margin
                    start_x = self.pdf.l_margin
                
                # Manual word wrapping
                words = safe_text.split(' ')
                current_line = ""
                
                for word in words:
                    test_line = current_line + (" " if current_line else "") + word
                    test_width = self.pdf.get_string_width(test_line)
                    
                    if test_width <= available_width or not current_line:
                        current_line = test_line
                    else:
                        # Write current line and start new one
                        try:
                            self.pdf.write(5, current_line)
                        except:
                            self.pdf.write(5, current_line.encode('ascii', 'replace').decode('ascii'))
                        
                        self.pdf.ln()
                        
                        # Set proper indentation for continuation
                        if self.in_list_item:
                            current_indent = self.get_current_indent()
                            bullet_width = 15
                            self.pdf.set_x(current_indent + bullet_width)
                            self.list_item_first_line = False
                        else:
                            self.pdf.set_x(self.pdf.l_margin)
                        
                        current_line = word
                
                # Write remaining text
                if current_line:
                    try:
                        self.pdf.write(5, current_line + " ")
                    except:
                        self.pdf.write(5, current_line.encode('ascii', 'replace').decode('ascii') + " ")
                
                self.current_text = ""
                self.last_was_block_element = False
                if self.in_list_item:
                    self.list_item_first_line = False
    
    # Create the PDF and parse the HTML
    pdf = FPDF()
    pdf.add_page()
    
    # Add Unicode font support
    try:
        # Try to add a Unicode font (DejaVu Sans supports most Unicode characters)
        pdf.add_font('DejaVu', '', 'DejaVuSans.ttf', uni=True)
        pdf.add_font('DejaVu', 'B', 'DejaVuSans-Bold.ttf', uni=True)
        pdf.add_font('DejaVu', 'I', 'DejaVuSans-Oblique.ttf', uni=True)
        default_font = 'DejaVu'
        unicode_support = True
    except:
        # Fallback to Arial if Unicode fonts are not available
        default_font = 'Arial'
        unicode_support = False
    
    pdf.set_font(default_font, size=12)
    
    parser = HTMLToPDFParser(pdf)
    parser.unicode_support = unicode_support
    parser.feed(html_text)
    parser.flush_text()  # Flush any remaining text
    
    # Save to BytesIO and return
    pdf_fp = BytesIO()
    pdf.output(pdf_fp)
    pdf_fp.seek(0)
    return pdf_fp.getvalue()

def process_pdf_inline_styles(pdf, line, indent, unicode_support=False, default_font='Arial'):
    """Helper to process a line with inline markdown for PDF `write` method, handling wrapping."""
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|\`.*?\`)', line)
    for part in parts:
        if not part: continue
        
        style = ''
        font = default_font
        size = 12
        if part.startswith('**') and part.endswith('**'):
            content = part[2:-2]
            style = 'B'
        elif part.startswith('*') and part.endswith('*'):
            content = part[1:-1]
            style = 'I'
        elif part.startswith('`') and part.endswith('`'):
            content = part[1:-1]
            font = 'Courier' if not unicode_support else default_font
            style = ''
            size = 10
        else:
            content = part
        
        pdf.set_font(font, style, size)
        
        # Handle Unicode encoding
        safe_content = content if unicode_support else content.encode('latin-1', 'replace').decode('latin-1')
        
        # Manual word wrapping for the `write` method
        words = safe_content.split(' ')
        for word in words:
            word_to_write = word + ' '
            word_width = pdf.get_string_width(word_to_write)
            if pdf.get_x() + word_width > pdf.w - pdf.r_margin:
                pdf.ln()
                pdf.set_x(indent)
            pdf.write(5, word_to_write)
    
    pdf.set_font(default_font, '', 12)



# Page Configuration
st.set_page_config(layout="wide", page_title="RoboGarden Instructor Copilot")

# --- Initialize Session State ---
if 'generated_course_text' not in st.session_state:
    st.session_state.generated_course_text = ""
if 'generated_quiz_text' not in st.session_state:
    st.session_state.generated_quiz_text = ""


# --- Styling (Inspired by the uploaded images) ---
def load_css():
    """
    Loads custom CSS to style the Streamlit application according to the RoboGarden theme.
    Colors are extracted from the provided banner.jpg and colors.png.
    - Canary Yellow (#ffc300)
    - Blue (#3f7cac)
    - Green (#8bc53f)
    - Red (#e53238)
    - Light Background (#ecffd6)
    """
    css = """
    <style>
        /* Import a playful, rounded font that matches the banner's style */
        @import url('https://fonts.googleapis.com/css2?family=Fredoka+One&display=swap');

        /* Main app background */
        .stApp {
            background-color: #ecffd6; /* A clean, light background */
        }
        
        /* Main title styling */
        .main h1:first-of-type {
            font-family: 'Fredoka One', cursive;
            color: #8bc53f; /* Robo-Green from logo */
        }

        /* Other titles and headers */
        h1:not(.main h1:first-of-type), h2, h3 {
            font-family: 'Fredoka One', cursive;
            color: #3f7cac; /* Robo-Blue from logo */
        }

        /* Buttons */
        .stButton>button {
            font-family: 'Fredoka One', cursive;
            border: 2px solid #ffc300; /* Canary-Yellow from logo */
            border-radius: 25px;
            color: #ffffff;
            background-color: #ffc300; /* Canary-Yellow from logo */
            padding: 12px 28px;
            font-size: 16px;
            font-weight: bold;
            text-transform: uppercase;
            transition: all 0.3s ease-in-out;
            box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        }
        .stButton>button:hover {
            transform: translateY(-2px);
            background-color: #8bc53f; /* Robo-Green from logo */
            color: white;
            border-color: #8bc53f;
            box-shadow: 0 6px 8px rgba(0,0,0,0.15);
        }

        /* Quiz Button Styling - Robo-Blue with Canary-Yellow hover */
        .quiz-button>button {
            font-family: 'Fredoka One', cursive;
            border: 2px solid #3f7cac; /* Robo-Blue from logo */
            border-radius: 25px;
            color: #ffffff;
            background-color: #3f7cac; /* Robo-Blue from logo */
            padding: 12px 28px;
            font-size: 16px;
            font-weight: bold;
            text-transform: uppercase;
            transition: all 0.3s ease-in-out;
            box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        }
        .quiz-button>button:hover {
            transform: translateY(-2px);
            background-color: #ffc300; /* Canary-Yellow from logo */
            color: white;
            border-color: #ffc300;
            box-shadow: 0 6px 8px rgba(0,0,0,0.15);
        }

        /* Download Buttons Styling */
        .stDownloadButton>button {
            font-family: 'Fredoka One', cursive;
            border: 2px solid #8bc53f; /* Robo-Green from logo */
            border-radius: 25px;
            color: #ffffff;
            background-color: #8bc53f; /* Robo-Green from logo */
            padding: 12px 28px;
            font-size: 14px;
            font-weight: bold;
            text-transform: uppercase;
            transition: all 0.3s ease-in-out;
            box-shadow: 0 4px 6px rgba(0,0,0,0.1);
            width: 100%;
        }
        
        .stDownloadButton>button:hover {
            transform: translateY(-2px);
            background-color: #6fa832; /* Darker green for hover */
            color: white;
            border-color: #6fa832;
            box-shadow: 0 6px 8px rgba(0,0,0,0.15);
        }
        .stDownloadButton>button:active {
            transform: translateY(0px);
            background-color: #5a8a28; /* Even darker green for click */
            border-color: #5a8a28;
            box-shadow: 0 2px 4px rgba(0,0,0,0.2);
        }

        /* Tabs styling */
        .stTabs [data-baseweb="tab-list"] {
            gap: 2px; /* Reduce gap between tabs */
            border-bottom: 3px solid #3f7cac; /* Blue underline for the tab bar */
        }
        .stTabs [data-baseweb="tab"] {
            height: 50px;
            white-space: pre-wrap;
            background-color: #eaf4fc; /* Lighter Blue Accent */
            border-radius: 8px 8px 0px 0px;
            gap: 10px;
            padding: 10px 20px;
            color: #3f7cac !important;
            font-family: 'Fredoka One', cursive;
            margin: 0;
        }
        .stTabs [aria-selected="true"] {
            background-color: #3f7cac; /* Robo-Blue */
            color: white !important;
            font-weight: bold;
        }
        
        /* Sidebar/Column styling for controls */
        div[data-testid="stVerticalBlock"] > [data-testid="stVerticalBlock"] > [data-testid="stVerticalBlock"] {
             background-color: #fff9e6; /* Light Yellow to match buttons */
             border-radius: 10px;
             padding: 20px;
        }

        /* Main content area styling */
        .main .block-container {
            padding-top: 1rem; /* Reduced padding */
        }
        
        /* Success/Info/Warning boxes */
        .stAlert {
            border-radius: 10px;
            border-width: 2px;
        }
        
        /* Style for success messages to use the green color */
        .stAlert[data-baseweb="notification-positive"] {
            border-color: #8bc53f;
        }
        
        /* Style for warning messages to use the red color */
        .stAlert[data-baseweb="notification-negative"] {
             border-color: #e53238;
        }

    </style>
    """
    st.markdown(css, unsafe_allow_html=True)

load_css()
if not DOWNLOAD_ENABLED:
    st.warning("Please install `fpdf2` and `python-docx` to enable download functionality (`pip install fpdf2 python-docx`).")


try:
    st.image("static/images/banner2.png", use_container_width=True)
except Exception:
    st.info("Banner image not found. Please add banner2.png to the static/images directory.")

# --- App Header ---
st.title("RoboGarden Instructor Co-pilot🤖")
st.write("Your creative partner for building amazing educational content!")

# --- API Key Check and Model Initialization ---
try:
    # Use a dummy key if not found in secrets for graceful failure
    api_key = st.secrets.get("GOOGLE_API_KEY", "DUMMY_KEY_FOR_UI_DISPLAY")
    if api_key == "DUMMY_KEY_FOR_UI_DISPLAY":
        st.warning("API key not found. Please set your GOOGLE_API_KEY in Streamlit secrets. App is in read-only mode.")
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel('gemini-1.5-flash-latest')
except Exception as e:
    st.error(f"Failed to configure Google AI: {e}")
    st.stop()


# --- Main Application Tabs ---
tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "### 🤖 Course Architect", 
    "### 🧐 Content Reviewer", 
    "### 🚀 Future-Proofing Engine",
    "### 🛠️ Quiz Creator",
    "### 📚 Full Course Generator"
])

# --- TAB 1: Course Architect ---
with tab1:
    st.header("Build a New Course TOC from Scratch")
    st.markdown("Upload your raw materials (like lecture notes, articles, or a textbook chapter) and let the AI build a structured course for you.")
    
    col1, col2 = st.columns([1, 2]) # [sidebar_width, main_content_width]
    
    with col1:
        st.subheader("Blueprint 📝")
        
        uploaded_files_gen = st.file_uploader(
            "Upload your content here", 
            accept_multiple_files=True,
            type=['pdf', 'docx', 'txt'],
            key="gen_uploader"
        )
        
        course_length = st.selectbox("Course Length", ["Quick (Overview)", "Moderate (Standard)", "Detailed (In-depth)"], key="gen_length")
        target_audience = st.selectbox("Target Audience", ["High School Students", "Undergraduate Students", "Industry Professionals", "General Public"], key="gen_audience")
        course_tone = st.selectbox("Tone of Voice", ["Formal & Academic", "Conversational & Friendly", "Technical & Precise"], key="gen_tone")
        
        if st.button("✨ Generate Course ✨", use_container_width=True, key="gen_button"):
            if uploaded_files_gen:
                with col2:
                    with st.spinner("Analyzing documents and designing your course... This is where the magic happens! 🪄"):
                        try:
                            raw_text = extract_text_from_files(uploaded_files_gen)
                            if raw_text.strip():
                                prompt = prompt_utils.create_generation_prompt(raw_text, course_length, target_audience, course_tone)
                                response = model.generate_content(prompt, request_options={'timeout': 600})
                                st.session_state.generated_course_text = response.text # Save to session state
                                st.session_state.generated_quiz_text = "" # Clear any existing quiz
                            else:
                                st.warning("HMM! Could not extract text from the uploaded files. Please check the files and try again.")
                                st.session_state.generated_course_text = ""
                                st.session_state.generated_quiz_text = ""
                        except Exception as e:
                            st.error(f"An error occurred during generation: {e}")
                            st.session_state.generated_course_text = ""
                            st.session_state.generated_quiz_text = ""
            else:
                st.warning("Please upload at least one document to start building.")

        # Quiz Creator Panel - Only show if course is generated
        if st.session_state.generated_course_text:
            st.markdown("---")
            st.subheader("Quiz Creator 🧠")
            
            quiz_type = st.radio("Quiz Type:", 
                               ["Multiple Choice", "True/False", "Short Answer", "Mixed"], 
                               key="course_quiz_type")
            
            quiz_difficulty = st.radio("Difficulty Level:",
                                     ["Easy", "Medium", "Hard", "Mixed"], 
                                     key="course_quiz_difficulty")
            
            # Custom styling for quiz button
            st.markdown('<div class="quiz-button">', unsafe_allow_html=True)
            if st.button("🧠 Generate Quiz 🧠", use_container_width=True, key="course_quiz_button"):
                with col2:
                    with st.spinner("Crafting your quiz from the generated course... 📝"):
                        try:
                            prompt = prompt_utils.create_quiz_creator_prompt(
                                st.session_state.generated_course_text, 
                                difficulty_level=quiz_difficulty, 
                                question_type=quiz_type
                            )
                            response = model.generate_content(prompt, request_options={'timeout': 600})
                            st.session_state.generated_quiz_text = response.text
                            print(st.session_state.generated_quiz_text)
                        except Exception as e:
                            st.error(f"An error occurred during quiz creation: {e}")
                            st.session_state.generated_quiz_text = ""
            st.markdown('</div>', unsafe_allow_html=True)                        
    
    with col2:
        if not st.session_state.generated_course_text:
            st.info("Your generated course will appear here once you click the generate button.")
        else:
            st.success("YEAAH! Your course is ready!")
            st.markdown(st.session_state.generated_course_text) # Display from session state
            
            st.markdown("---") # Separator

            # Add download buttons if libraries are installed
            if DOWNLOAD_ENABLED:
                user_params = {
                    'audience': target_audience,
                    'length': course_length,
                }
                
                dl_col_1, dl_col_2, dl_col_3 = st.columns([2,2,2])
                with dl_col_1:
                    pdf_bytes = create_styled_pdf(st.session_state.generated_course_text)
                    course_filename = create_descriptive_filename(
                        "Course", 
                        user_params, 
                        st.session_state.generated_course_text, 
                        "pdf"
                    )
                    st.download_button(
                        label="⬇️ Download as PDF",
                        data=pdf_bytes,
                        file_name=course_filename,
                        mime="application/pdf",
                        use_container_width=True
                    )
                
                with dl_col_2:
                    docx_bytes = create_styled_docx(st.session_state.generated_course_text)
                    course_docx_filename = create_descriptive_filename(
                        "Course", 
                        user_params, 
                        st.session_state.generated_course_text, 
                        "docx"
                    )
                    st.download_button(
                        label="⬇️ Download as DOCX",
                        data=docx_bytes,
                        file_name=course_docx_filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                        
            # Button to clear the state and start over
            if st.button("Generate New Course", use_container_width=False):
                st.session_state.generated_course_text = ""
                st.session_state.generated_quiz_text = ""
                st.rerun()

            # Display generated quiz if available
            if st.session_state.generated_quiz_text:
                st.markdown("---")
                st.markdown(st.session_state.generated_quiz_text, unsafe_allow_html=True)
                
            # Quiz download buttons
            if DOWNLOAD_ENABLED:
                quiz_params = {
                    'audience': target_audience,
                    'quiz_type': quiz_type,
                    'difficulty': quiz_difficulty
                }
                
                quiz_dl_col_1, quiz_dl_col_2, quiz_dl_col_3 = st.columns([2,2,1])
                with quiz_dl_col_1:
                    quiz_pdf_bytes = create_styled_pdf(st.session_state.generated_quiz_text)
                    quiz_filename = create_descriptive_filename(
                        "Quiz", 
                        quiz_params, 
                        st.session_state.generated_quiz_text, 
                        "pdf"
                    )
                    st.download_button(
                        label="⬇️ PDF",
                        data=quiz_pdf_bytes,
                        file_name=quiz_filename,
                        mime="application/pdf",
                        use_container_width=True
                    )
                with quiz_dl_col_2:
                    quiz_docx_bytes = create_styled_docx(st.session_state.generated_quiz_text)
                    quiz_docx_filename = create_descriptive_filename(
                        "Quiz", 
                        quiz_params, 
                        st.session_state.generated_quiz_text, 
                        "docx"
                    )
                    st.download_button(
                        label="⬇️ DOCX",
                        data=quiz_docx_bytes,
                        file_name=quiz_docx_filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )


# --- TAB 2: Content Reviewer ---
with tab2:
    st.header("Validate and Improve an Existing Course")
    st.markdown("Upload a complete course document. The AI will act as a pedagogical expert, checking for clarity, engagement, and style issues.")
    
    uploaded_file_val = st.file_uploader(
        "Upload your course document here", 
        accept_multiple_files=False, 
        type=['pdf', 'docx', 'txt'],
        key="val_uploader"
    )

    if st.button("🔍 Validate Course 🔍", use_container_width=True, key="val_button"):
        if uploaded_file_val:
            with st.spinner("Our expert is proofreading your course... 🧐"):
                try:
                    raw_text = extract_text_from_files([uploaded_file_val])
                    if raw_text.strip():
                        prompt = prompt_utils.create_validation_prompt(raw_text)
                        response = model.generate_content(prompt, request_options={'timeout': 600})
                        st.success("YEAAH! Validation complete.")
                        st.markdown(response.text)
                    else:
                        st.warning("HMM! Could not extract text from the uploaded file.")
                except Exception as e:
                    st.error(f"An error occurred during validation: {e}")
        else:
            st.warning("Please upload a document to validate.")

# --- TAB 3: Future-Proofing Engine ---
with tab3:
    st.header("Update a Course with the Latest Trends")
    st.markdown("Is your content still current? Upload a course, and the AI will scan the horizon for new trends, technologies, and research, suggesting relevant updates.")

    uploaded_file_upd = st.file_uploader(
        "Upload your course document to update", 
        accept_multiple_files=False, 
        type=['pdf', 'docx', 'txt'],
        key="upd_uploader"
    )

    if st.button("🚀 Future-Proof Course 🚀", use_container_width=True, key="upd_button"):
        if uploaded_file_upd:
            with st.spinner("Scanning the future for updates... 📡"):
                try:
                    raw_text = extract_text_from_files([uploaded_file_upd])
                    if raw_text.strip():
                        prompt = prompt_utils.create_updater_prompt(raw_text)
                        response = model.generate_content(prompt, request_options={'timeout': 600})
                        st.success("YEAAH! Here are the suggested updates.")
                        st.markdown(response.text)
                    else:
                        st.warning("HMM! Could not extract text from the uploaded file.")
                except Exception as e:
                    st.error(f"An error occurred during the update process: {e}")
        else:
            st.warning("Please upload a document to update.")

# --- TAB 4: Quiz Creator ---
with tab4:
    st.header("Create a Quiz for Your Course")
    st.markdown("Upload your course content, and the AI will generate a comprehensive quiz to test understanding and retention.")
    st.radio("Select the type of quiz you want to create:", 
             ["Multiple Choice", "True/False", "Short Answer","Mixed"], 
             key="quiz_type")
    
    st.radio("Select the difficulty level of the quiz:",
             ["Easy", "Medium", "Hard", "Mixed"], 
             key="quiz_difficulty")
    
    uploaded_file_quiz = st.file_uploader(
        "Upload your course document for quiz creation", 
        accept_multiple_files=False, 
        type=['pdf', 'docx', 'txt'],
        key="quiz_uploader"
    )

    if st.button("🛠️ Create Quiz 🛠️", use_container_width=True, key="quiz_button"):
        if uploaded_file_quiz:
            with st.spinner("Crafting your quiz... 📝"):
                try:
                    raw_text = extract_text_from_files([uploaded_file_quiz])
                    if raw_text.strip():
                        prompt = prompt_utils.create_quiz_creator_prompt(raw_text,difficulty_level=st.session_state.quiz_difficulty, question_type=st.session_state.quiz_type)
                        response = model.generate_content(prompt, request_options={'timeout': 600})
                        st.success("YEAAH! Your quiz is ready!")
                        st.markdown(response.text)
                    else:
                        st.warning("HMM! Could not extract text from the uploaded file.")
                except Exception as e:
                    st.error(f"An error occurred during quiz creation: {e}")
        else:
            st.warning("Please upload a document to create a quiz.")

with tab5:
    st.header("Build a New Course from Scratch")
    st.markdown("Upload your raw materials (like lecture notes, articles, or a textbook chapter) and let the AI build a structured course for you.")
    
    col1, col2 = st.columns([1, 2]) # [sidebar_width, main_content_width]
    
    with col1:
        st.subheader("Blueprint 📝")
        
        uploaded_files_gen = st.file_uploader(
            "Upload your content here", 
            accept_multiple_files=True,
            type=['pdf', 'docx', 'txt'],
        )

        course_length = st.selectbox("Course Length", ["Quick (Overview)", "Moderate (Standard)", "Detailed (In-depth)"],key="gen_length_5")
        target_audience = st.selectbox("Target Audience", ["High School Students", "Undergraduate Students", "Industry Professionals", "General Public"], key="gen_audience_5")
        course_tone = st.selectbox("Tone of Voice", ["Formal & Academic", "Conversational & Friendly", "Technical & Precise"], key="gen_tone_5")
        if st.button("✨ Generate Full Course Content ✨", use_container_width=True, key="gen_button_2"):
            if uploaded_files_gen:
                with col2:
                    with st.spinner("Analyzing documents and designing your full course... This may take a few minutes 🪄"):
                        try:
                            raw_text = extract_text_from_files(uploaded_files_gen)
                            if raw_text.strip():
                                # Initialize container for full course content
                                full_course_content = ""
                                
                                # Store content sections separately
                                if 'course_sections' not in st.session_state:
                                    st.session_state.course_sections = {}
                                
                                # Step 1: Generate Table of Contents
                                st.info("Step 1/3: Generating table of contents...")
                                toc_prompt = f"""
                                Based on the following materials, create a detailed table of contents for a {course_length} course targeted at {target_audience} with a {course_tone} tone.
                                Create a well-structured course with 3-5 main lessons. Include lesson titles and 3-5 subtopics for each lesson.
                                Format as a proper Markdown table of contents with # for main lesson titles and ## for subtopics.
                                
                                Materials:
                                {raw_text[:10000]}
                                """
                                toc_response = model.generate_content(toc_prompt, request_options={'timeout': 300})
                                table_of_contents = toc_response.text
                                full_course_content += "# Course Table of Contents\n\n" + table_of_contents + "\n\n"
                                
                                # Store the table of contents section
                                st.session_state.course_sections["Table of Contents"] = table_of_contents
                                
                                # Extract lesson titles
                                lesson_titles = []
                                for line in table_of_contents.split('\n'):
                                    if line.strip().startswith('# '):
                                        lesson_titles.append(line.strip()[2:])
                                
                                # Step 2: Generate each lesson content with its quiz
                                st.info(f"Step 2/3: Generating {len(lesson_titles)} lessons with quizzes...")
                                progress_bar = st.progress(0)
                                
                                for idx, title in enumerate(lesson_titles):
                                    # Generate lesson content
                                    lesson_prompt = f"""
                                    Create detailed content for the lesson titled "{title}" for a {course_length} course targeted at {target_audience}.
                                    Use a {course_tone} tone. Include explanations, examples, and key concepts.
                                    Format using Markdown with proper headings, bullet points, and emphasis where appropriate.
                                    Base the content on these materials:
                                    
                                    {raw_text[:15000]}
                                    """
                                    lesson_response = model.generate_content(lesson_prompt, request_options={'timeout': 300})
                                    lesson_content = lesson_response.text
                                    full_course_content += f"\n\n# {title}\n\n{lesson_content}\n\n"
                                    
                                    # Store this lesson's content
                                    st.session_state.course_sections[f"Lesson {idx+1}: {title}"] = lesson_content
                                    
                                    # Immediately generate quiz for this lesson
                                    st.info(f"Creating quiz for lesson: {title}")
                                    quiz_prompt = prompt_utils.create_quiz_creator_prompt(
                                        lesson_content,
                                        difficulty_level="Medium", 
                                        question_type="Mixed"
                                    )
                                    quiz_response = model.generate_content(quiz_prompt, request_options={'timeout': 300})
                                    quiz_content = quiz_response.text
                                    full_course_content += f"\n\n## Quiz: {title}\n\n{quiz_content}\n\n"
                                    
                                    # Store this quiz's content
                                    st.session_state.course_sections[f"Quiz {idx+1}: {title}"] = quiz_content
                                    
                                    # Update progress bar
                                    progress_bar.progress((idx + 1) / len(lesson_titles))
                                
                                # Add "Full Course" option which contains everything
                                st.session_state.course_sections["Full Course"] = full_course_content
                                
                                # Save to session state
                                st.session_state.generated_course_text = full_course_content
                                st.session_state.course_generated = True
                                
                                # Display success message
                                st.success("✅ Full course generation complete with table of contents, lessons, and quizzes!")
                                
                            else:
                                st.warning("Could not extract text from the uploaded files. Please check the files and try again.")
                        except Exception as e:
                            st.error(f"An error occurred during generation: {e}")
            else:
                st.warning("Please upload at least one document to start building your course.")
    
    with col2:
        if 'course_generated' in st.session_state and st.session_state.course_generated:
            # Create a dropdown to select what content to view
            section_options = list(st.session_state.course_sections.keys())
            selected_section = st.selectbox("Choose what to view:", section_options)
            
            # Display the selected section
            st.markdown(st.session_state.course_sections[selected_section])
            
            # Add download buttons
            user_params = {
                'audience': st.session_state.get('gen_audience_2', 'General'),
                'length': st.session_state.get('gen_length_2', 'Standard'),
            }
            
            # Determine if we're downloading the full course or a section
            download_content = st.session_state.course_sections[selected_section]
            section_name = selected_section.replace(" ", "_").replace(":", "")
            
            dl_col_1, dl_col_2 = st.columns(2)
            with dl_col_1:
                pdf_bytes = create_styled_pdf(download_content)
                course_filename = create_descriptive_filename(
                    section_name, 
                    user_params, 
                    download_content, 
                    "pdf"
                )
                st.download_button(
                    label=f"⬇️ Download {selected_section} as PDF",
                    data=pdf_bytes,
                    file_name=course_filename,
                    mime="application/pdf",
                    use_container_width=True
                )
            
            with dl_col_2:
                docx_bytes = create_styled_docx(download_content)
                course_docx_filename = create_descriptive_filename(
                    section_name, 
                    user_params, 
                    download_content, 
                    "docx"
                )
                st.download_button(
                    label=f"⬇️ Download {selected_section} as DOCX",
                    data=docx_bytes,
                    file_name=course_docx_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )