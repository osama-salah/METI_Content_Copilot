"""
File generation utilities for creating PDF, DOCX, and HTML files with proper formatting.
This module handles the conversion of markdown/HTML content to various downloadable formats.
"""

import re
from io import BytesIO
from datetime import datetime
import streamlit as st

# --- Optional dependencies for download functionality ---
try:
    from fpdf import FPDF
    import docx
    from docx.shared import Pt
    DOWNLOAD_ENABLED = True
except ImportError:
    DOWNLOAD_ENABLED = False


def create_descriptive_filename(content_type, user_params=None, content_text="", file_extension="pdf"):
    """Create descriptive filenames based on content and user selections"""
    
    # Extract course topic/title if possible
    topic = "Course"
    if content_text:
        lines = content_text.split('\n')
        for line in lines:
            if line.strip().startswith('# '):
                topic = line.strip()[2:].strip()
                # Clean for filename
                topic = re.sub(r'[^\w\s-]', '', topic)
                topic = re.sub(r'[-\s]+', '_', topic)[:30]
                break
    
    # Build filename parts
    parts = [topic]
    
    if user_params:
        if 'audience' in user_params:
            audience_short = user_params['audience'].split()[0]  # "High", "Undergraduate", etc.
            parts.append(audience_short)
        
        if 'length' in user_params:
            length_short = user_params['length'].split()[0]  # "Quick", "Moderate", "Detailed"
            parts.append(length_short)
            
        if 'quiz_type' in user_params:
            quiz_type_short = user_params['quiz_type'].replace(' ', '')
            parts.append(quiz_type_short)
    
    parts.append(content_type)
    
    # Add timestamp for uniqueness
    timestamp = datetime.now().strftime("%m%d_%H%M")
    parts.append(timestamp)
    
    filename = "_".join(parts) + f".{file_extension}"
    return filename


def create_styled_docx(text):
    """Generates a DOCX file from a Markdown string with styling for headers, lists, and code."""
    doc = docx.Document()
    in_code_block = False
    code_block_started = False  # Track if we just started a code block
    
    for line in text.split('\n'):
        # Code block handling - FIX: Check for triple backticks
        if line.strip().startswith('```'):
            in_code_block = not in_code_block
            if in_code_block:
                # Only add spacing before code block if the previous element wasn't empty
                code_block_started = True
            else:
                # Add some spacing after code block
                doc.add_paragraph()
                code_block_started = False
            continue

        if in_code_block:
            # For code blocks, add paragraph with shaded background
            p = doc.add_paragraph(line)
            p.style = 'No Spacing'
            
            # Set font for code
            if p.runs:
                font = p.runs[0].font
            else:
                run = p.add_run('')
                font = run.font
            font.name = 'Courier New'
            font.size = Pt(10)
            
            # Add shaded background to the paragraph
            from docx.oxml.shared import qn
            from docx.oxml import parse_xml
            
            # Create shading element with light gray background
            shading_elm = parse_xml(r'<w:shd {} w:fill="F0F0F0"/>'.format(
                'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            ))
            p._element.get_or_add_pPr().append(shading_elm)
            
            # Add border around code block paragraphs
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
            # Set paragraph spacing for code blocks
            paragraph_format = p.paragraph_format
            paragraph_format.left_indent = Pt(12)  # Slight indentation
            paragraph_format.right_indent = Pt(12)
            
            # Only add space before the first line of the code block
            if code_block_started:
                paragraph_format.space_before = Pt(6)
                code_block_started = False
            else:
                paragraph_format.space_before = Pt(2)
            
            paragraph_format.space_after = Pt(2)
            
            continue

        # Other markdown handling (rest of the function remains the same)
        line = line.strip()
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith(('* ', '- ')):
            p = doc.add_paragraph(style='List Bullet')
            # Handle inline styles (bold, italic, code)
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|\`.*?\`)', line[2:])
            for part in parts:
                if not part: continue
                if part.startswith('**') and part.endswith('**'):
                    p.add_run(part[2:-2]).bold = True
                elif part.startswith('*') and part.endswith('*'):
                    p.add_run(part[1:-1]).italic = True
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Courier New'
                    # Add light gray background for inline code
                    from docx.oxml.shared import qn
                    from docx.oxml import parse_xml
                    shading_elm = parse_xml(r'<w:shd {} w:fill="E8E8E8"/>'.format(
                        'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
                    ))
                    run._element.get_or_add_rPr().append(shading_elm)
                else:
                    p.add_run(part)
        elif line:
            p = doc.add_paragraph()
            # Handle inline styles (bold, italic, code)
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|\`.*?\`)', line)
            for part in parts:
                if not part: continue
                if part.startswith('**') and part.endswith('**'):
                    p.add_run(part[2:-2]).bold = True
                elif part.startswith('*') and part.endswith('*'):
                    p.add_run(part[1:-1]).italic = True
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Courier New'
                    # Add light gray background for inline code
                    from docx.oxml.shared import qn
                    from docx.oxml import parse_xml
                    shading_elm = parse_xml(r'<w:shd {} w:fill="E8E8E8"/>'.format(
                        'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
                    ))
                    run._element.get_or_add_rPr().append(shading_elm)
                else:
                    p.add_run(part)

    doc_fp = BytesIO()
    doc.save(doc_fp)
    doc_fp.seek(0)
    return doc_fp.getvalue()


def create_html_docx(html_text):
    """
    Generates a DOCX file from an HTML string with styling 
    for headers, lists, bold, italic, and other HTML elements.
    """
    from html.parser import HTMLParser
    import html
    
    class HTMLToDocxParser(HTMLParser):
        def __init__(self, doc):
            super().__init__()
            self.doc = doc
            self.current_paragraph = None
            self.current_run = None
            self.style_stack = []  # Stack to track nested styles
            self.in_list = False
            self.list_stack = []  # Stack to track nested lists
            self.list_counters = []  # Stack to track list item counters
            self.current_text = ""
            self.pending_paragraph = False
            
        def handle_starttag(self, tag, attrs):
            self.flush_text()
            
            if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(tag[1])
                # Add spacing before headers (except if it's the first element)
                if len(self.doc.paragraphs) > 0:
                    self.doc.add_paragraph()
                self.current_paragraph = self.doc.add_heading('', level=level)
                self.current_run = None
            elif tag == 'p':
                # Only add paragraph if we don't already have one or if we have content
                if self.current_paragraph is None or len(self.doc.paragraphs) > 0:
                    self.current_paragraph = self.doc.add_paragraph()
                self.current_run = None
            elif tag in ['b', 'strong']:
                self.style_stack.append('bold')
                self._ensure_run()
            elif tag in ['i', 'em']:
                self.style_stack.append('italic')
                self._ensure_run()
            elif tag == 'code':
                self.style_stack.append('code')
                self._ensure_run()
            elif tag == 'pre':
                # Add spacing before pre block
                if len(self.doc.paragraphs) > 0:
                    self.doc.add_paragraph()
                self.current_paragraph = self.doc.add_paragraph()
                self.current_paragraph.style = 'No Spacing'
                self.style_stack.append('pre')
                self._ensure_run()
            elif tag == 'ul':
                self.in_list = True
                self.list_stack.append('bullet')
                self.list_counters.append(0)  # Not used for bullets but keeps stack aligned
                # Add spacing before list
                if len(self.doc.paragraphs) > 0 and not self.current_paragraph:
                    self.doc.add_paragraph()
            elif tag == 'ol':
                self.in_list = True
                self.list_stack.append('number')
                self.list_counters.append(0)  # Counter for numbered items
                # Add spacing before list
                if len(self.doc.paragraphs) > 0 and not self.current_paragraph:
                    self.doc.add_paragraph()
            elif tag == 'li':
                if self.list_stack:
                    list_type = self.list_stack[-1]
                    if list_type == 'number':
                        # Increment counter for this list level
                        self.list_counters[-1] += 1
                        self.current_paragraph = self.doc.add_paragraph(style='List Number')
                    else:
                        self.current_paragraph = self.doc.add_paragraph(style='List Bullet')
                else:
                    # Fallback to bullet if no list context
                    self.current_paragraph = self.doc.add_paragraph(style='List Bullet')
                self.current_run = None
            elif tag == 'br':
                if self.current_paragraph is None:
                    self.current_paragraph = self.doc.add_paragraph()
                self.current_paragraph.add_run().add_break()
                
        def handle_endtag(self, tag):
            self.flush_text()
            
            if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                self.current_paragraph = None
                self.current_run = None
            elif tag == 'p':
                self.current_paragraph = None
                self.current_run = None
            elif tag in ['b', 'strong', 'i', 'em', 'code']:
                if self.style_stack and self.style_stack[-1] in ['bold', 'italic', 'code']:
                    self.style_stack.pop()
                self.current_run = None  # Force new run creation
            elif tag == 'pre':
                if 'pre' in self.style_stack:
                    self.style_stack.remove('pre')
                self.current_paragraph = None
                self.current_run = None
                # Add spacing after pre block
                self.doc.add_paragraph()
            elif tag in ['ul', 'ol']:
                if self.list_stack:
                    self.list_stack.pop()
                    self.list_counters.pop()
                self.in_list = len(self.list_stack) > 0
                if not self.in_list:
                    # Add spacing after outermost list
                    self.doc.add_paragraph()
            elif tag == 'li':
                self.current_paragraph = None
                self.current_run = None
                
        def handle_data(self, data):
            # Decode HTML entities and clean up excessive whitespace
            data = html.unescape(data)
            # Only add non-empty text
            if data.strip():
                self.current_text += data.strip() + " "
            
        def flush_text(self):
            if self.current_text.strip():
                self._ensure_paragraph()
                self._ensure_run()
                
                # Apply styles from stack
                text_to_add = self.current_text.strip()
                run = self.current_run.add_run(text_to_add) if hasattr(self.current_run, 'add_run') else self.current_paragraph.add_run(text_to_add)
                
                # Apply formatting based on style stack
                if 'bold' in self.style_stack:
                    run.bold = True
                if 'italic' in self.style_stack:
                    run.italic = True
                if 'code' in self.style_stack or 'pre' in self.style_stack:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                    # Add light gray background for code
                    from docx.oxml.shared import qn
                    from docx.oxml import parse_xml
                    try:
                        shading_elm = parse_xml(r'<w:shd {} w:fill="E8E8E8"/>'.format(
                            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
                        ))
                        run._element.get_or_add_rPr().append(shading_elm)
                    except:
                        pass  # Skip styling if it fails
                
                self.current_text = ""
        
        def _ensure_paragraph(self):
            if self.current_paragraph is None:
                self.current_paragraph = self.doc.add_paragraph()
        
        def _ensure_run(self):
            if self.current_run is None:
                self._ensure_paragraph()
                self.current_run = self.current_paragraph
    
    # Create the document and parse the HTML
    doc = docx.Document()
    parser = HTMLToDocxParser(doc)
    parser.feed(html_text)
    parser.flush_text()  # Flush any remaining text
    
    # Save to BytesIO and return
    doc_fp = BytesIO()
    doc.save(doc_fp)
    doc_fp.seek(0)
    return doc_fp.getvalue()


def process_pdf_inline_styles_safe(pdf, line, indent):
    """Simplified version of inline style processing with better error handling"""
    try:
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|\`.*?\`)', line)
        for part in parts:
            if not part: 
                continue
            
            style = ''
            font = 'Helvetica'
            size = 12
            
            if part.startswith('**') and part.endswith('**'):
                content = part[2:-2]
                style = 'B'
            elif part.startswith('*') and part.endswith('*'):
                content = part[1:-1]
                style = 'I'
            elif part.startswith('`') and part.endswith('`'):
                content = part[1:-1]
                font = 'Courier'
                size = 10
            else:
                content = part
            
            # Ensure safe encoding
            safe_content = content.encode('latin-1', 'replace').decode('latin-1')
            
            pdf.set_font(font, style, size)
            
            # Simple word wrapping
            words = safe_content.split(' ')
            for word in words:
                word_to_write = word + ' '
                word_width = pdf.get_string_width(word_to_write)
                if pdf.get_x() + word_width > pdf.w - pdf.r_margin:
                    pdf.ln()
                    pdf.set_x(indent)
                pdf.write(5, word_to_write)
        
        pdf.set_font('Helvetica', '', 12)
        
    except Exception as e:
        # Fallback to simple text if styling fails
        safe_text = line.encode('latin-1', 'replace').decode('latin-1')
        pdf.set_font('Helvetica', '', 12)
        pdf.multi_cell(pdf.w - 2 * pdf.l_margin, 5, safe_text)


def create_fallback_pdf(text):
    """Create a simple PDF if the styled version fails"""
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font('Helvetica', size=12)
        
        # Simple text conversion - remove markdown
        clean_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)  # Remove bold
        clean_text = re.sub(r'\*([^*]+)\*', r'\1', clean_text)  # Remove italic
        clean_text = re.sub(r'`([^`]+)`', r'\1', clean_text)  # Remove code
        clean_text = re.sub(r'^#+\s*', '', clean_text, flags=re.MULTILINE)  # Remove headers
        
        # Encode safely
        safe_text = clean_text.encode('latin-1', 'replace').decode('latin-1')
        
        pdf.multi_cell(pdf.w - 2 * pdf.l_margin, 5, safe_text)
        
        pdf_fp = BytesIO()
        pdf.output(pdf_fp)
        pdf_fp.seek(0)
        return pdf_fp.getvalue()
        
    except Exception as e:
        st.error(f"Even fallback PDF generation failed: {e}")
        return b"PDF generation failed"


def create_styled_pdf(text):
    """
    Generates a PDF file from a Markdown string with native styling
    for headers, lists, bold, italic, and code.
    """
    try:
        pdf = FPDF()
        pdf.add_page()
        
        # Simplified font handling - stick to built-in fonts
        pdf.set_font('Helvetica', size=12)
        unicode_support = False  # Disable Unicode to avoid font issues
        default_font = 'Helvetica'
        
        effective_page_width = pdf.w - 2 * pdf.l_margin

        in_code_block = False
        for line in text.split('\n'):
            # Code block handling
            if line.strip().startswith('```'):
                in_code_block = not in_code_block
                if in_code_block:
                    pdf.set_font('Courier', size=10)
                    pdf.set_fill_color(240, 240, 240)
                    pdf.ln(2)
                else:
                    pdf.set_font('Helvetica', size=12)
                    pdf.set_fill_color(255, 255, 255)
                    pdf.ln(5)
                continue

            if in_code_block:
                pdf.set_x(pdf.l_margin)
                # Ensure safe encoding for code blocks
                safe_line = line.encode('latin-1', 'replace').decode('latin-1')
                pdf.multi_cell(effective_page_width, 5, safe_line, border=0, fill=True)
                continue

            # Other markdown handling
            line = line.strip()
            if not line:
                pdf.ln(5)
                continue

            # Safe text encoding for all content
            if line.startswith('# '):
                pdf.set_font('Helvetica', 'B', 16)
                safe_text = line[2:].strip().encode('latin-1', 'replace').decode('latin-1')
                pdf.multi_cell(effective_page_width, 8, safe_text)
                pdf.ln(4)
            elif line.startswith('## '):
                pdf.set_font('Helvetica', 'B', 14)
                safe_text = line[3:].strip().encode('latin-1', 'replace').decode('latin-1')
                pdf.multi_cell(effective_page_width, 7, safe_text)
                pdf.ln(3)
            elif line.startswith('### '):
                pdf.set_font('Helvetica', 'B', 12)
                safe_text = line[4:].strip().encode('latin-1', 'replace').decode('latin-1')
                pdf.multi_cell(effective_page_width, 6, safe_text)
                pdf.ln(2)
            elif line.startswith(('* ', '- ')):
                pdf.set_x(pdf.l_margin)
                initial_x = pdf.get_x()
                pdf.write(5, '- ')
                process_pdf_inline_styles_safe(pdf, line[2:].strip(), initial_x + 5)
                pdf.ln()
            else:
                pdf.set_x(pdf.l_margin)
                process_pdf_inline_styles_safe(pdf, line, pdf.l_margin)
                pdf.ln()

        pdf_fp = BytesIO()
        pdf.output(pdf_fp)
        pdf_fp.seek(0)
        return pdf_fp.getvalue()
        
    except Exception as e:
        # If PDF generation fails, create a simple fallback PDF
        st.error(f"PDF generation failed: {e}")
        return create_fallback_pdf(text)


def create_html_pdf(html_text):
    """
    Generates a PDF file from an HTML string with styling
    for headers, lists, bold, italic, and other HTML elements.
    Properly handles ordered lists with numbers and nested lists with letters.
    """
    from html.parser import HTMLParser
    import html
    
    class HTMLToPDFParser(HTMLParser):
        def __init__(self, pdf):
            super().__init__()
            self.pdf = pdf
            self.font_stack = [('Helvetica', '', 12)]  # Stack to track font changes
            self.in_code = False
            self.in_pre = False
            self.list_stack = []  # Stack to track list types and levels
            self.list_counters = []  # Stack to track counters for each list level
            self.list_indent_level = 0
            self.base_indent = 10  # Base indentation for lists
            self.current_text = ""
            self.unicode_support = False
            self.last_was_block_element = False
            self.in_list_item = False
            self.list_item_first_line = True
            
        def get_current_indent(self):
            """Calculate current indentation based on list nesting level"""
            return self.pdf.l_margin + (self.list_indent_level * self.base_indent)
            
        def handle_starttag(self, tag, attrs):
            self.flush_text()
            
            if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                # Add spacing before headers
                if not self.last_was_block_element and self.pdf.get_y() > 20:
                    self.pdf.ln(8)
                
                level = int(tag[1])
                size = max(16 - (level - 1) * 2, 10)
                self.pdf.set_font('Helvetica', 'B', size)
                self.last_was_block_element = True
                
            elif tag == 'p':
                # Add spacing before paragraphs (but not if we're in a list item)
                if not self.last_was_block_element and self.pdf.get_y() > 20 and not self.in_list_item:
                    self.pdf.ln(4)
                self.last_was_block_element = True
                
            elif tag == 'b' or tag == 'strong':
                current_font = self.font_stack[-1]
                new_style = current_font[1] + 'B' if 'B' not in current_font[1] else current_font[1]
                self.font_stack.append((current_font[0], new_style, current_font[2]))
                self.pdf.set_font(current_font[0], new_style, current_font[2])
                
            elif tag == 'i' or tag == 'em':
                current_font = self.font_stack[-1]
                new_style = current_font[1] + 'I' if 'I' not in current_font[1] else current_font[1]
                self.font_stack.append((current_font[0], new_style, current_font[2]))
                self.pdf.set_font(current_font[0], new_style, current_font[2])
                
            elif tag == 'code':
                self.in_code = True
                self.font_stack.append(('Courier', '', 10))
                self.pdf.set_font('Courier', '', 10)
                
            elif tag == 'pre':
                # Add spacing before pre block
                if not self.last_was_block_element:
                    self.pdf.ln(4)
                self.in_pre = True
                self.pdf.set_font('Courier', '', 10)
                self.pdf.set_fill_color(240, 240, 240)
                self.last_was_block_element = True
                
            elif tag == 'ul':
                # Add spacing before lists (only for outermost lists)
                if self.list_indent_level == 0 and not self.last_was_block_element:
                    self.pdf.ln(4)
                
                self.list_stack.append('ul')
                self.list_counters.append(0)  # Not used for bullets but keeps stack aligned
                self.list_indent_level += 1
                self.last_was_block_element = True
                
            elif tag == 'ol':
                # Add spacing before lists (only for outermost lists)
                if self.list_indent_level == 0 and not self.last_was_block_element:
                    self.pdf.ln(4)
                
                self.list_stack.append('ol')
                self.list_counters.append(0)  # Counter for numbered items
                self.list_indent_level += 1
                self.last_was_block_element = True
                
            elif tag == 'li':
                self.in_list_item = True
                self.list_item_first_line = True
                
                # Move to new line for list item
                if self.pdf.get_x() > self.pdf.l_margin:
                    self.pdf.ln()
                
                # Set proper indentation
                current_indent = self.get_current_indent()
                self.pdf.set_x(current_indent)
                
                if self.list_stack:
                    list_type = self.list_stack[-1]
                    if list_type == 'ol':
                        # Increment counter and use number
                        self.list_counters[-1] += 1
                        counter = self.list_counters[-1]
                        
                        # Format number based on nesting level
                        if self.list_indent_level == 1:
                            # Top level: 1. 2. 3.
                            number_text = f"{counter}. "
                        elif self.list_indent_level == 2:
                            # Second level: a) b) c)
                            letter = chr(ord('a') + counter - 1) if counter <= 26 else f"item{counter}"
                            number_text = f"{letter}) "
                        elif self.list_indent_level == 3:
                            # Third level: i. ii. iii.
                            roman_numerals = ['i', 'ii', 'iii', 'iv', 'v', 'vi', 'vii', 'viii', 'ix', 'x']
                            roman = roman_numerals[counter - 1] if counter <= len(roman_numerals) else f"item{counter}"
                            number_text = f"{roman}. "
                        else:
                            # Deeper levels: fallback to numbers
                            number_text = f"{counter}. "
                            
                        try:
                            self.pdf.write(5, number_text)
                        except:
                            # Fallback for encoding issues
                            safe_text = number_text.encode('ascii', 'replace').decode('ascii')
                            self.pdf.write(5, safe_text)
                    else:
                        # Use different bullet styles based on nesting level
                        if self.list_indent_level == 1:
                            bullet_char = '• ' if self.unicode_support else '* '
                        elif self.list_indent_level == 2:
                            bullet_char = '◦ ' if self.unicode_support else '- '
                        else:
                            bullet_char = '▪ ' if self.unicode_support else '+ '
                            
                        try:
                            self.pdf.write(5, bullet_char)
                        except:
                            self.pdf.write(5, '- ')
                else:
                    # Fallback bullet if no list context
                    try:
                        self.pdf.write(5, '• ')
                    except:
                        self.pdf.write(5, '- ')
                    
            elif tag == 'br':
                self.pdf.ln()
                # Maintain indentation if we're in a list item
                if self.in_list_item:
                    # Calculate the indent for continuation lines (after the bullet/number)
                    current_indent = self.get_current_indent()
                    bullet_width = 15  # Approximate width of bullet/number + space
                    self.pdf.set_x(current_indent + bullet_width)
                    self.list_item_first_line = False
                else:
                    self.pdf.set_x(self.pdf.l_margin)
                
        def handle_endtag(self, tag):
            self.flush_text()
            
            if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                self.pdf.ln(6)  # Space after headers
                self.pdf.set_font('Helvetica', '', 12)  # Reset font
                self.pdf.set_x(self.pdf.l_margin)  # Reset position
                self.last_was_block_element = True
                
            elif tag == 'p':
                # Only add line break if we're not in a list item
                if not self.in_list_item:
                    self.pdf.ln(4)  # Space after paragraphs
                    self.pdf.set_x(self.pdf.l_margin)  # Reset position
                self.last_was_block_element = True
                
            elif tag in ['b', 'strong', 'i', 'em', 'code']:
                if len(self.font_stack) > 1:
                    self.font_stack.pop()
                current_font = self.font_stack[-1]
                self.pdf.set_font(current_font[0], current_font[1], current_font[2])
                if tag == 'code':
                    self.in_code = False
                    
            elif tag == 'pre':
                self.in_pre = False
                self.pdf.set_fill_color(255, 255, 255)  # Reset background
                self.pdf.set_font('Helvetica', '', 12)  # Reset font
                self.pdf.ln(6)  # Space after pre block
                self.pdf.set_x(self.pdf.l_margin)  # Reset position
                self.last_was_block_element = True
                
            elif tag in ['ul', 'ol']:
                if self.list_stack:
                    self.list_stack.pop()
                    self.list_counters.pop()
                    self.list_indent_level -= 1
                    
                # Add spacing after outermost list
                if self.list_indent_level == 0:
                    self.pdf.ln(4)
                    self.pdf.set_x(self.pdf.l_margin)
                    self.last_was_block_element = True
                
            elif tag == 'li':
                self.in_list_item = False
                self.list_item_first_line = True
                self.pdf.ln()  # New line after list item
                
        def handle_data(self, data):
            """Handle text data between HTML tags"""
            # Decode HTML entities and clean up excessive whitespace
            data = html.unescape(data)
            # Only add non-empty text
            if data.strip():
                self.current_text += data.strip() + " "
        
        def flush_text(self):
            """Write accumulated text to PDF with proper wrapping and indentation"""
            if self.current_text.strip():
                # Handle Unicode encoding
                safe_text = self.current_text.strip()
                if not self.unicode_support:
                    safe_text = safe_text.encode('latin-1', 'replace').decode('latin-1')
                
                # Calculate available width for text
                if self.in_list_item:
                    current_indent = self.get_current_indent()
                    if self.list_item_first_line:
                        # First line: text starts after bullet/number
                        bullet_width = 15
                        available_width = self.pdf.w - current_indent - bullet_width - self.pdf.r_margin
                        start_x = current_indent + bullet_width
                    else:
                        # Continuation lines: align with first line text
                        bullet_width = 15
                        available_width = self.pdf.w - current_indent - bullet_width - self.pdf.r_margin
                        start_x = current_indent + bullet_width
                        self.pdf.set_x(start_x)
                else:
                    available_width = self.pdf.w - self.pdf.l_margin - self.pdf.r_margin
                    start_x = self.pdf.l_margin
                
                # Manual word wrapping
                words = safe_text.split(' ')
                current_line = ""
                
                for word in words:
                    test_line = current_line + (" " if current_line else "") + word
                    test_width = self.pdf.get_string_width(test_line)
                    
                    if test_width <= available_width or not current_line:
                        current_line = test_line
                    else:
                        # Write current line and start new one
                        try:
                            self.pdf.write(5, current_line)
                        except:
                            self.pdf.write(5, current_line.encode('ascii', 'replace').decode('ascii'))
                        
                        self.pdf.ln()
                        
                        # Set proper indentation for continuation
                        if self.in_list_item:
                            current_indent = self.get_current_indent()
                            bullet_width = 15
                            self.pdf.set_x(current_indent + bullet_width)
                            self.list_item_first_line = False
                        else:
                            self.pdf.set_x(self.pdf.l_margin)
                        
                        current_line = word
                
                # Write remaining text
                if current_line:
                    try:
                        self.pdf.write(5, current_line + " ")
                    except:
                        self.pdf.write(5, current_line.encode('ascii', 'replace').decode('ascii') + " ")
                
                self.current_text = ""
                self.last_was_block_element = False
                if self.in_list_item:
                    self.list_item_first_line = False
    
    # Create the PDF and parse the HTML
    pdf = FPDF()
    pdf.add_page()
    
    # Add Unicode font support
    try:
        # Try to add a Unicode font (DejaVu Sans supports most Unicode characters)
        pdf.add_font('DejaVu', '', 'DejaVuSans.ttf', uni=True)
        pdf.add_font('DejaVu', 'B', 'DejaVuSans-Bold.ttf', uni=True)
        pdf.add_font('DejaVu', 'I', 'DejaVuSans-Oblique.ttf', uni=True)
        default_font = 'DejaVu'
        unicode_support = True
    except:
        # Fallback to Helvetica if Unicode fonts are not available
        default_font = 'Helvetica'
        unicode_support = False
    
    pdf.set_font(default_font, size=12)
    
    parser = HTMLToPDFParser(pdf)
    parser.unicode_support = unicode_support
    parser.feed(html_text)
    parser.flush_text()  # Flush any remaining text
    
    # Save to BytesIO and return
    pdf_fp = BytesIO()
    pdf.output(pdf_fp)
    pdf_fp.seek(0)
    return pdf_fp.getvalue()
